#!/usr/bin/env python3
"""Build a human-reviewable DOCX packet from a stratified random sample of
the 959-case v2 flip candidate set.

Sampling follows the same spirit as ../build_variability_sample.py
(stratify by boundary_id, fixed seed 42 for reproducibility) but uses a
floor of 1 row/family rather than 2: at n=50 across 33 families, a floor of
2 would force 66 rows (33 x 2), overshooting "about 50" by a lot. A floor of
1 guarantees every family is represented at least once, and the remaining
slots are drawn uniformly at random from the leftover pool so larger
families get a proportionally higher chance at a second pick.
"""

from __future__ import annotations

import argparse
import csv
import random
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HERE = Path(__file__).resolve().parent
STUDY_ROOT = HERE.parent


def stratified_sample_floor1(rows: list[dict], n: int, seed: int) -> list[dict]:
    by_family: dict[str, list[dict]] = defaultdict(list)
    for row in rows:
        by_family[row["boundary_id"]].append(row)

    rng = random.Random(seed)
    families = sorted(by_family)

    picked: list[dict] = []
    remaining_by_family: dict[str, list[dict]] = {}
    for fam in families:
        pool = by_family[fam][:]
        rng.shuffle(pool)
        picked.append(pool[0])
        remaining_by_family[fam] = pool[1:]

    leftover = [row for fam in families for row in remaining_by_family[fam]]
    rng.shuffle(leftover)
    extra_needed = max(0, n - len(picked))
    picked.extend(leftover[:extra_needed])

    return picked

FIELD_LABELS = [
    ("mechanism", "Mechanism"),
    ("initial_category", "Initial category"),
    ("initial_subcategory", "Initial subcategory"),
    ("final_category", "Final category (after disclosure)"),
    ("final_subcategory", "Final subcategory (after disclosure)"),
    ("hidden_fact", "Hidden fact"),
    ("fact_as_answer", "Fact as answer (what user says when asked)"),
    ("relevant_question_topic", "Relevant follow-up question topic"),
    ("counterfactual_category", "Counterfactual category"),
    ("counterfactual_subcategory", "Counterfactual subcategory"),
    ("counterfactual_hidden_fact", "Counterfactual hidden fact"),
    ("counterfactual_fact_as_answer", "Counterfactual fact as answer"),
    ("safety_sensitive", "Safety sensitive"),
    ("vetting_note", "Vetting note"),
]


def add_field_table(document: Document, row: dict) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = (Pt(150), Pt(330))
    for key, label in FIELD_LABELS:
        value = (row.get(key) or "").strip()
        if not value:
            continue
        cells = table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = value
        for cell, width in zip(cells, widths):
            cell.width = width


def add_reviewer_block(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Reviewer sign-off:")
    run.bold = True
    document.add_paragraph("Valid?  ☐ Yes   ☐ No   ☐ Unsure")
    document.add_paragraph("Notes: " + "_" * 70)
    document.add_paragraph("_" * 90)


def build_document(sample: list[dict], pool_total: int, seed: int, n: int) -> Document:
    document = Document()

    title = document.add_heading("FETCH Flip Benchmark — Manual Validity Review Sample", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    families_represented = len({r["boundary_id"] for r in sample})
    meta = document.add_paragraph()
    meta.add_run(f"Generated: {date.today().isoformat()}\n").italic = True
    meta.add_run(
        f"Sample: {len(sample)} of {pool_total} scenarios "
        f"({families_represented} of 33 boundary families represented)\n"
    )
    meta.add_run(
        f"Sampling method: stratified by boundary_id (1 guaranteed row per family, remaining "
        f"slots drawn uniformly at random from the leftover pool), random seed {seed}.\n"
    )

    intro = document.add_paragraph(
        "Each scenario below is a candidate case for the FETCH classification-flip benchmark: an "
        "opening query that is genuinely ambiguous between an “initial” category/subcategory and "
        "a “final” category/subcategory, resolved only once the user discloses a hidden fact. "
        "A counterfactual variant (different hidden fact, different resolution) is also included so the "
        "same opening query can test both directions. When reviewing each scenario, check: (1) does the "
        "opening query read as something a real person would plausibly say, and is it actually ambiguous "
        "between the initial and final categories without more information; (2) does the hidden fact "
        "plausibly and specifically resolve that ambiguity toward the final category; (3) does the "
        "counterfactual hidden fact plausibly resolve it toward the counterfactual category instead; "
        "(4) for safety-sensitive scenarios, is the framing appropriate and non-gratuitous."
    )
    intro.paragraph_format.space_after = Pt(12)

    document.add_page_break()

    by_family: dict[str, list[dict]] = defaultdict(list)
    for row in sample:
        by_family[row["boundary_id"]].append(row)

    for family in sorted(by_family):
        family_rows = sorted(by_family[family], key=lambda r: r["scenario_id"])
        document.add_heading(f"{family}  (n={len(family_rows)})", level=1)

        for row in family_rows:
            heading_text = f"{row['scenario_id']}  —  {row['direction']}  —  {row['flip_type']}"
            document.add_heading(heading_text, level=2)

            query_p = document.add_paragraph()
            query_p.paragraph_format.left_indent = Pt(18)
            query_run = query_p.add_run(f"“{row['opening_query'].strip()}”")
            query_run.italic = True

            flip_p = document.add_paragraph()
            flip_p.add_run("Flip tested: ").bold = True
            flip_p.add_run(
                f"{row['initial_category']} > {row['initial_subcategory']}  →  "
                f"{row['final_category']} > {row['final_subcategory']}"
            )

            if (row.get("safety_sensitive") or "").strip().lower() == "true":
                flag_p = document.add_paragraph()
                flag_run = flag_p.add_run("⚠ SAFETY SENSITIVE")
                flag_run.bold = True
                flag_run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

            add_field_table(document, row)
            add_reviewer_block(document)
            document.add_paragraph()

    return document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", default=str(STUDY_ROOT / "candidates/flip_candidates_v2.csv"))
    parser.add_argument("--out-docx", default=str(HERE / "flip_sample_review.docx"))
    parser.add_argument("--out-ids", default=str(HERE / "sample_scenario_ids.csv"))
    parser.add_argument("--n", type=int, default=50)
    parser.add_argument("--seed", type=int, default=42)
    args = parser.parse_args()

    with open(args.source, newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        rows = list(reader)

    sample = stratified_sample_floor1(rows, args.n, args.seed)
    sample.sort(key=lambda r: (r["boundary_id"], r["scenario_id"]))

    with open(args.out_ids, "w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["scenario_id", "boundary_id", "direction", "flip_type", "safety_sensitive"])
        for row in sample:
            writer.writerow(
                [row["scenario_id"], row["boundary_id"], row["direction"], row["flip_type"], row["safety_sensitive"]]
            )

    document = build_document(sample, pool_total=len(rows), seed=args.seed, n=args.n)
    document.save(args.out_docx)

    by_family = defaultdict(int)
    safety_count = 0
    for row in sample:
        by_family[row["boundary_id"]] += 1
        if (row.get("safety_sensitive") or "").strip().lower() == "true":
            safety_count += 1

    print(f"Sampled {len(sample)} of {len(rows)} scenarios across {len(by_family)} of 33 families (seed={args.seed})")
    print(f"Safety-sensitive scenarios in sample: {safety_count}")
    print(f"Wrote {args.out_docx}")
    print(f"Wrote {args.out_ids}")
    for fam in sorted(by_family):
        print(f"  {fam}: {by_family[fam]}")


if __name__ == "__main__":
    main()
